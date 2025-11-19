from __future__ import annotations
import logging
from typing import Optional
from io import BytesIO
from docx import Document

from .document_parser import DocumentParser
from ..exceptions import ExtractionError

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    name = "python-docx"

    def supports(self, *, mime: Optional[str], ext: Optional[str]) -> bool:
        docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if mime is not None:
            return mime == docx_mime
        return ext is not None and ext.lower() == ".docx"

    def extract_text(self, content: bytes) -> str:
        try:
            doc = Document(BytesIO(content))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as e:
            raise ExtractionError(f"Failed to extract text from DOCX: {e}") from e
